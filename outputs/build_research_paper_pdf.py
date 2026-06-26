from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"C:\Users\jiten\btp_final2")
OUT_DIR = ROOT / "outputs" / "final_artifacts"
ASSET_DIR = OUT_DIR / "assets"
RESULTS_DIR = ROOT / "TSFFM-Depression-Detection" / "results"
PDF_PATH = OUT_DIR / "TSFFM_Depression_Detection_Research_Paper.pdf"
DOCX_PATH = OUT_DIR / "TSFFM_Depression_Detection_Research_Paper.docx"


TITLE = "Visual Landmark-Based Two-Stream Temporal Fusion for AI-Assisted Depression Screening from Interview Video"
AUTHORS = "Jitendra Choudhary (22BEC059), Kabir Singh Khair (22BEC061)"
AFFILIATION = "Department of Electronics and Communication Engineering, PDPM IIITDM Jabalpur, India"
SUPERVISOR = "Supervisor: Dr. Amit Vishwakarma"


ABSTRACT = (
    "Depression screening commonly depends on clinical interviews, questionnaires, and professional "
    "interpretation. These methods remain necessary, but access constraints and under-reporting motivate "
    "assistive tools that can analyze behavioral evidence in a repeatable manner. This paper presents a "
    "final-year project prototype for AI-assisted depression screening from interview video using a "
    "landmark-based Two-Stream Feature Fusion Model (TSFFM). The implemented system extracts a 68-point "
    "facial landmark subset using MediaPipe FaceMesh and shoulder-level pose landmarks using MediaPipe "
    "Pose. The features are normalized, sampled to a fixed 360-frame sequence, projected through separate "
    "face and body streams, fused into a 160-dimensional temporal representation, smoothed by a "
    "one-dimensional temporal convolution, and classified by a bidirectional LSTM with temporal attention. "
    "The model is integrated into a FastAPI backend and React/Vite frontend that supports video upload, "
    "probability display, detection-quality reporting, and PDF report generation. On the available "
    "validation evaluation, the selected checkpoint achieved 57.67% accuracy, 0.4688 precision, 0.4255 "
    "recall, 0.4461 F1-score, and 0.6597 AUC-ROC over 352 segments. A recall-oriented candidate achieved "
    "0.8227 recall and 0.5873 F1-score but lower accuracy and AUC. The results indicate that the prototype "
    "learns measurable signal but is not clinically ready."
)


KEYWORDS = (
    "Depression screening; behavioral computing; facial landmarks; pose landmarks; temporal attention; "
    "LSTM; multimodal fusion; DAIC-WOZ; FastAPI."
)


SECTIONS = [
    (
        "I. INTRODUCTION",
        [
            "Depression is a major public-health concern. The World Health Organization reports that approximately 332 million people worldwide have depression, with prevalence varying across age and demographic groups. Screening and diagnosis remain clinical tasks, typically conducted through interview, questionnaire, and professional judgment. However, access limitations, stigma, and self-report bias can delay care. Automated behavioral analysis can therefore be useful as an assistive screening component, provided that its limitations are explicitly stated and clinical decision-making remains with qualified professionals.",
            "Video interviews contain non-verbal signals that may be relevant to depression assessment, including facial expressiveness, gaze behavior, posture, and psychomotor activity. Prior depression-recognition research has frequently used the DAIC-WOZ and AVEC challenge settings, where audio, video, and transcript modalities are collected during human-computer or semi-clinical interviews. These datasets support reproducible experimentation, but they are small relative to modern deep-learning requirements, and generalization across cohorts remains difficult.",
            "The project reported in this paper began as a mid-semester TSFFM concept centered on two visual streams: a facial stream and a body-posture stream. The final implementation changes the early frame-level CNN concept into a compact landmark-based temporal architecture. Specifically, the deployed pipeline extracts face and pose coordinate features, normalizes them, uses separate multilayer streams for face and body, fuses the streams, and uses bidirectional temporal modeling with attention.",
        ],
    ),
    (
        "II. RELATED WORK",
        [
            "The Distress Analysis Interview Corpus contains clinical interviews designed to support the diagnosis of psychological distress conditions such as depression, anxiety, and post-traumatic stress disorder. AVEC 2017 provided benchmark conditions for affect and depression analysis using common data and evaluation procedures. These benchmark settings motivate this project because they frame depression detection as a multimodal behavioral-computing problem rather than as a conventional image-classification task.",
            "MediaPipe FaceMesh is related to real-time facial surface geometry estimation from monocular video, where dense face meshes are inferred efficiently on mobile hardware. MediaPipe Pose and BlazePose-style systems provide body landmarks suitable for real-time human pose estimation. The current prototype uses these tools as deterministic feature extractors rather than training an image-to-depression network directly.",
            "Depression-related behavior is temporal. LSTM networks address long-range temporal dependency learning in recurrent neural networks. This project uses a bidirectional LSTM to model the fixed-length landmark sequence. Class imbalance is addressed using class weights, weighted sampling, and focal loss.",
        ],
    ),
    (
        "III. SYSTEM ARCHITECTURE",
        [
            "The implemented system consists of a React/Vite dashboard, a FastAPI prediction API, a MediaPipe-based feature extractor, a PyTorch TSFFM temporal model, training/evaluation scripts, and a ReportLab-based PDF report generator. The frontend handles video upload, patient/session metadata, loading states, probability display, and report download. The backend validates uploads, saves the temporary file, extracts features, runs inference, generates a report, returns a structured response, and deletes the temporary upload.",
            "The deployed data flow begins with an uploaded video file. The backend samples frames at 5 FPS, extracts facial and shoulder pose landmarks, pads or truncates the sequence to 360 frames, normalizes coordinates per frame, flattens the face sequence to shape 360 x 272, flattens the pose sequence to shape 360 x 8, and converts both arrays into tensors with batch dimension. The model produces logits for two classes: not depressed and depressed.",
            "The current safety mechanisms are software and communication safeguards, not clinical safeguards. File extensions are checked before processing. Uploaded files are deleted in a cleanup block after inference. The generated report contains a medical disclaimer. Quantitative evidence is unavailable for adversarial upload handling, privacy compliance, subgroup fairness, model calibration, or clinical review.",
        ],
    ),
    (
        "IV. METHODOLOGY",
        [
            "For each sampled frame, FaceMesh produces dense facial landmarks. The implementation selects 68 indices to form a standard facial subset. Each facial point stores x, y, z, and confidence, yielding 272 values per frame. Pose extraction uses the left and right shoulders from MediaPipe Pose; each point stores x, y, z, and visibility, yielding 8 values per frame.",
            "Each sequence is centered by subtracting the per-frame mean over spatial coordinates. It is scaled by the maximum per-frame distance from the center, with zero-distance protection. Videos shorter than the target length are padded by repeating the last available frame; longer videos are truncated to 360 sampled frames. If no readable frames exist, zero arrays are returned.",
            "The face stream maps each 272-dimensional vector to a 128-dimensional latent representation. The body stream maps each 8-dimensional pose vector to a 32-dimensional representation. The features are concatenated into a 160-dimensional temporal sequence. A one-dimensional convolution with kernel size 3 is applied over time. A bidirectional LSTM maps the fused sequence into hidden states, and temporal attention forms a context vector before binary classification.",
        ],
    ),
    (
        "V. EXPERIMENTAL ANALYSIS",
        [
            "The evaluation script loads the validation split, restores the trained checkpoint, performs inference, and computes accuracy, precision, recall, F1-score, and AUC-ROC. It also saves a text classification report, confusion matrix, ROC curve, and training-history plots.",
            "The selected checkpoint has stronger AUC-ROC and accuracy, while the recall-oriented candidate detects a larger fraction of depressed samples at the cost of ranking quality and overall accuracy. In a screening setting, recall is important because false negatives are harmful. However, a high-recall operating point must be calibrated with clinical input to manage false-positive burden.",
            "The selected checkpoint classification report includes 211 not-depressed validation segments and 141 depressed validation segments. The not-depressed class achieved 0.64 precision, 0.68 recall, and 0.66 F1-score. The depressed class achieved 0.47 precision, 0.43 recall, and 0.45 F1-score. The lower depressed-class recall is a major concern for an assistive screening system.",
        ],
    ),
    (
        "VI. DISCUSSION",
        [
            "The final prototype demonstrates a functioning end-to-end visual screening system. Its main engineering improvement over the mid-semester version is the transition from a conceptual two-stream slide model to a deployed temporal pipeline with feature extraction, model inference, frontend interaction, probability display, and report generation. The main research improvement is the addition of bidirectional temporal modeling and attention.",
            "The quantitative results do not support a clinical deployment claim. The selected checkpoint's AUC-ROC of 0.6597 suggests that the probability scores contain some discriminative information, but the depressed-class recall of 0.4255 is not acceptable for a high-sensitivity screening use case without further calibration and validation.",
            "The deployed body stream is narrower than the original mid-semester concept. The mid-semester deck described 33 body landmarks and a 132-dimensional pose vector; the final implementation uses only two shoulder landmarks and an 8-dimensional pose vector. This verified implementation difference should be described transparently.",
        ],
    ),
    (
        "VII. ETHICS AND LIMITATIONS",
        [
            "This system processes interview video, which is sensitive behavioral and biometric data. Any deployment outside a controlled academic demonstration would require informed consent, secure storage, retention limits, access control, and ethics approval. The generated reports must not be treated as medical diagnoses.",
            "The main limitations are validation scope, incomplete dataset provenance, limited pose representation, absence of ablation experiments, absence of repeated-run statistics, absence of calibration analysis, no latency benchmark, and no clinical validation. The current evaluation should be treated as prototype evidence rather than clinical evidence.",
            "A model-version consistency issue also exists: the current evaluation source instantiates TSFFM_LSTM with one LSTM layer, whereas training and inference instantiate the model with two LSTM layers. The saved metric files are reported as repository artifacts, but exact reproduction requires aligning these scripts.",
        ],
    ),
    (
        "VIII. FUTURE WORK AND CONCLUSION",
        [
            "The most important next step is to improve validation quality before increasing model complexity. A publishable extension should include participant-level split documentation, external test evaluation, repeated runs with confidence intervals, and ablations for face stream, pose stream, fusion, temporal convolution, LSTM, and attention.",
            "The model should also be extended to richer multimodal evidence. Audio prosody, pause duration, speech rate, and transcript semantics are strongly relevant to depression interviews and are used in prior DAIC-WOZ studies. A tri-stream architecture combining visual landmarks, acoustic descriptors, and transcript embeddings would better match the multimodal nature of clinical interviews.",
            "This paper presented an implementation-grounded TSFFM prototype for AI-assisted depression screening from interview video. The system extracts visual landmarks, fuses facial and pose streams, models temporal behavior with a bidirectional LSTM and attention, and exposes the model through a web application with report generation. The current evidence supports the claim that the prototype functions end to end and learns measurable validation signal. It does not support claims of clinical readiness, generalizable diagnostic performance, or deployment safety.",
        ],
    ),
]


LIT_TABLE = [
    ["Paper or Source", "Contribution", "Limitation", "Relevance"],
    ["DAIC corpus", "Clinical interviews with verbal and non-verbal distress indicators.", "Access and partition constraints limit direct replication.", "Defines the interview-based behavioral setting."],
    ["AVEC 2017", "Benchmark conditions for depression and affect recognition.", "Benchmarks do not replace independent validation.", "Provides evaluation context."],
    ["Zhang et al.", "Text, audio, and video depression detection on DAIC-WOZ.", "Uses modalities beyond this project.", "Motivates audio/text future work."],
    ["FaceMesh", "Real-time facial geometry from monocular video.", "Not designed for clinical inference.", "Supports landmark features."],
    ["BlazePose", "Real-time pose landmarks.", "Pose accuracy does not imply depression validity.", "Supports posture features."],
]


METRICS_TABLE = [
    ["Model", "Accuracy", "Precision", "Recall", "F1", "AUC"],
    ["Selected checkpoint", "57.67%", "0.4688", "0.4255", "0.4461", "0.6597"],
    ["Recall-oriented candidate", "53.69%", "0.4567", "0.8227", "0.5873", "0.5454"],
]


EVIDENCE_TABLE = [
    ["Experiment", "Evidence", "Result", "Limitation"],
    ["Selected validation", "classification_report.txt", "57.67% accuracy, 0.6597 AUC", "No external test set."],
    ["Recall candidate", "retrained_candidate report", "0.8227 recall, 0.5873 F1", "Lower AUC and accuracy."],
    ["Training dynamics", "history JSON and plots", "Overfitting pressure visible", "No repeated-run variance."],
    ["End-to-end app", "Frontend/backend source", "Upload, inference, report workflow", "No latency benchmark."],
]


REFERENCES = [
    "[1] World Health Organization, \"Depressive disorder (depression),\" Aug. 29, 2025. Available: https://www.who.int/news-room/fact-sheets/detail/depression.",
    "[2] J. Gratch et al., \"The Distress Analysis Interview Corpus of human and computer interviews,\" Proc. LREC, 2014.",
    "[3] USC Institute for Creative Technologies, \"DAIC-WOZ Depression Database: Documentation for AVEC 2017,\" 2017.",
    "[4] F. Ringeval et al., \"AVEC 2017: Real-life depression, and affect recognition workshop and challenge,\" Proc. AVEC, 2017.",
    "[5] W. Zhang, K. Mao, and J. Chen, \"A multimodal approach for detection and assessment of depression using text, audio and video,\" Phenomics, vol. 4, no. 3, pp. 234-249, 2024.",
    "[6] Y. Kartynnik et al., \"Real-time facial surface geometry from monocular video on mobile GPUs,\" arXiv:1907.06724, 2019.",
    "[7] I. Grishchenko et al., \"BlazePose GHUM Holistic: Real-time 3D human landmarks and pose estimation,\" arXiv:2206.11678, 2022.",
    "[8] S. Hochreiter and J. Schmidhuber, \"Long short-term memory,\" Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997.",
    "[9] T.-Y. Lin et al., \"Focal loss for dense object detection,\" Proc. IEEE ICCV, pp. 2980-2988, 2017.",
]


FIGURES = [
    ("system_pipeline.png", ASSET_DIR / "system_pipeline.png", "Fig. 1. End-to-end inference pipeline implemented in the project."),
    ("model_architecture.png", ASSET_DIR / "model_architecture.png", "Fig. 2. Implemented TSFFM temporal architecture."),
    ("metrics_summary.png", ASSET_DIR / "metrics_summary.png", "Fig. 3. Metric comparison between the selected checkpoint and recall-oriented candidate."),
    ("confusion_matrix.png", RESULTS_DIR / "confusion_matrix.png", "Fig. 4. Confusion matrix saved by the evaluation script."),
    ("roc_curve.png", RESULTS_DIR / "roc_curve.png", "Fig. 5. ROC curve saved by the evaluation script."),
    ("accuracy_history.png", RESULTS_DIR / "accuracy_history.png", "Fig. 6. Training and validation accuracy history."),
    ("loss_history.png", RESULTS_DIR / "loss_history.png", "Fig. 7. Training and validation loss history."),
]


def page_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Times-Roman", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(0.75 * inch, 0.45 * inch, "TSFFM Depression Detection Research Paper")
    canvas.drawRightString(7.75 * inch, 0.45 * inch, str(canvas.getPageNumber()))
    canvas.restoreState()


def para_style(name, parent=None, **kwargs):
    return ParagraphStyle(name, parent=parent, **kwargs)


def build_pdf():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    body = para_style("IEEEBody", styles["BodyText"], fontName="Times-Roman", fontSize=9.2, leading=10.7, alignment=TA_JUSTIFY, spaceAfter=4)
    heading = para_style("IEEEHeading", styles["Heading2"], fontName="Times-Bold", fontSize=10.2, leading=12, alignment=TA_LEFT, spaceBefore=8, spaceAfter=4)
    title = para_style("Title", styles["Title"], fontName="Times-Bold", fontSize=16, leading=18, alignment=TA_CENTER, spaceAfter=8)
    meta = para_style("Meta", styles["BodyText"], fontName="Times-Roman", fontSize=9.2, leading=11, alignment=TA_CENTER, spaceAfter=3)
    abstract_head = para_style("AbsHead", styles["BodyText"], fontName="Times-Bold", fontSize=9.2, leading=10.5, alignment=TA_LEFT, spaceAfter=2)
    caption = para_style("Caption", styles["BodyText"], fontName="Times-Italic", fontSize=7.7, leading=9, alignment=TA_CENTER, spaceBefore=2, spaceAfter=5)
    ref_style = para_style("References", styles["BodyText"], fontName="Times-Roman", fontSize=8, leading=9, alignment=TA_LEFT, spaceAfter=2)

    doc = BaseDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.62 * inch,
        rightMargin=0.62 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    gap = 0.22 * inch
    col_width = (doc.width - gap) / 2
    frame_left = Frame(doc.leftMargin, doc.bottomMargin, col_width, doc.height, id="left")
    frame_right = Frame(doc.leftMargin + col_width + gap, doc.bottomMargin, col_width, doc.height, id="right")
    doc.addPageTemplates([PageTemplate(id="TwoCol", frames=[frame_left, frame_right], onPage=page_footer)])

    story = []
    story.append(Paragraph(TITLE, title))
    story.append(Paragraph(AUTHORS, meta))
    story.append(Paragraph(AFFILIATION, meta))
    story.append(Paragraph(SUPERVISOR, meta))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Abstract", abstract_head))
    story.append(Paragraph(ABSTRACT, body))
    story.append(Paragraph("<b>Keywords:</b> " + KEYWORDS, body))

    story.append(Paragraph("LITERATURE REVIEW SUMMARY", heading))
    story.append(make_table(LIT_TABLE, [0.75 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch], font_size=6.0))

    for title_text, paragraphs in SECTIONS:
        story.append(Paragraph(title_text, heading))
        for p in paragraphs:
            story.append(Paragraph(p, body))
        if title_text == "III. SYSTEM ARCHITECTURE":
            add_figure(story, FIGURES[0], caption)
        if title_text == "IV. METHODOLOGY":
            add_figure(story, FIGURES[1], caption)
        if title_text == "V. EXPERIMENTAL ANALYSIS":
            story.append(make_table(METRICS_TABLE, [1.0 * inch, 0.55 * inch, 0.55 * inch, 0.55 * inch, 0.45 * inch, 0.45 * inch], font_size=6.5))
            story.append(make_table(EVIDENCE_TABLE, [0.8 * inch, 0.9 * inch, 0.9 * inch, 1.0 * inch], font_size=6.0))
            add_figure(story, FIGURES[2], caption)
            add_figure(story, FIGURES[3], caption)
            add_figure(story, FIGURES[4], caption)
            add_figure(story, FIGURES[5], caption)
            add_figure(story, FIGURES[6], caption)

    story.append(Paragraph("MEDICAL DISCLAIMER", heading))
    story.append(Paragraph("The system described here is an AI-assisted screening prototype. It is not a diagnostic device and must not be used for medical decisions without professional clinical validation and qualified mental-health oversight.", body))

    story.append(Paragraph("REFERENCES", heading))
    for ref in REFERENCES:
        story.append(Paragraph(ref, ref_style))

    doc.build(story)
    return PDF_PATH


def make_table(data, widths, font_size=7):
    table = Table(data, colWidths=widths, hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), "Times-Roman", font_size),
                ("FONT", (0, 0), (-1, 0), "Times-Bold", font_size),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#9AA6B2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]
        )
    )
    return table


def add_figure(story, fig, caption_style):
    _, path, text = fig
    if path.exists():
        story.append(Spacer(1, 3))
        story.append(Image(str(path), width=3.35 * inch, height=2.05 * inch, kind="proportional"))
        story.append(Paragraph(text, caption_style))


def set_docx_font(run, size=None, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_docx_para(doc, text, style=None, size=10, bold=False, align=None, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_docx_font(r, size=size, bold=bold, italic=italic)
    return p


def add_docx_table(doc, rows):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, val in enumerate(rows[0]):
        set_cell(table.rows[0].cells[i], val, bold=True)
        shade_cell(table.rows[0].cells[i], "E8EEF5")
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], val)
    doc.add_paragraph()


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_docx_font(r, size=8.5, bold=bold)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_docx_para(doc, TITLE, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_para(doc, AUTHORS, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_para(doc, AFFILIATION, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_para(doc, SUPERVISOR, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_docx_para(doc, "Abstract", size=10, bold=True)
    add_docx_para(doc, ABSTRACT, size=10)
    add_docx_para(doc, "Keywords: " + KEYWORDS, size=9, italic=True)
    add_docx_para(doc, "Literature Review Summary", size=12, bold=True)
    add_docx_table(doc, LIT_TABLE)

    for title_text, paragraphs in SECTIONS:
        add_docx_para(doc, title_text, size=12, bold=True)
        for p in paragraphs:
            add_docx_para(doc, p, size=10)
        if title_text == "III. SYSTEM ARCHITECTURE":
            add_docx_figure(doc, FIGURES[0])
        if title_text == "IV. METHODOLOGY":
            add_docx_figure(doc, FIGURES[1])
        if title_text == "V. EXPERIMENTAL ANALYSIS":
            add_docx_table(doc, METRICS_TABLE)
            add_docx_table(doc, EVIDENCE_TABLE)
            for fig in FIGURES[2:]:
                add_docx_figure(doc, fig)

    add_docx_para(doc, "Medical Disclaimer", size=12, bold=True)
    add_docx_para(doc, "The system described here is an AI-assisted screening prototype. It is not a diagnostic device and must not be used for medical decisions without professional clinical validation and qualified mental-health oversight.", size=10)
    add_docx_para(doc, "References", size=12, bold=True)
    for ref in REFERENCES:
        add_docx_para(doc, ref, size=8)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def add_docx_figure(doc, fig):
    _, path, caption = fig
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(5.8))
        add_docx_para(doc, caption, size=8, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    pdf = build_pdf()
    docx = build_docx()
    print(f"PDF: {pdf}")
    print(f"DOCX: {docx}")


if __name__ == "__main__":
    main()
